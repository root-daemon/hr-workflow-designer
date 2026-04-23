"""
Generates CASE_STUDY.docx — HR Workflow Designer solution document.
Run:  python3 generate_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Palette ───────────────────────────────────────────────────────────────────
INK       = RGBColor(0x0f, 0x11, 0x17)
SURFACE   = RGBColor(0x1a, 0x1d, 0x27)
CARD      = RGBColor(0x21, 0x25, 0x2f)
BORDER_C  = RGBColor(0x2e, 0x33, 0x47)
ACCENT    = RGBColor(0x63, 0x66, 0xf1)
ACCENT2   = RGBColor(0x81, 0x8c, 0xf8)
TEXT      = RGBColor(0xe2, 0xe4, 0xed)
MUTED     = RGBColor(0x8b, 0x8f, 0xa8)
WHITE     = RGBColor(0xff, 0xff, 0xff)
GREEN     = RGBColor(0x22, 0xc5, 0x5e)
AMBER     = RGBColor(0xf5, 0x9e, 0x0b)
ROSE      = RGBColor(0xf4, 0x3f, 0x5e)
BLUE      = RGBColor(0x3b, 0x82, 0xf6)
PURPLE    = RGBColor(0xa8, 0x55, 0xf7)
CODE_BG   = RGBColor(0x0a, 0x0c, 0x12)
GREEN_LT  = RGBColor(0x86, 0xef, 0xac)
AMBER_LT  = RGBColor(0xfc, 0xd3, 0x4d)
BLUE_LT   = RGBColor(0x93, 0xc5, 0xfd)
PURPLE_LT = RGBColor(0xd8, 0xb4, 0xfe)
ROSE_LT   = RGBColor(0xfd, 0xa4, 0xaf)

# ── Helpers ───────────────────────────────────────────────────────────────────

def hex_color(color: RGBColor) -> str:
    return '%02X%02X%02X' % (color[0], color[1], color[2])

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{hex_color(color)}')
    tcPr.append(shd)

def set_doc_bg(doc, color: RGBColor):
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), f'{hex_color(color)}')
    doc.element.insert(0, bg)
    settings = doc.settings.element
    disp = OxmlElement('w:displayBackgroundShape')
    settings.append(disp)

def set_para_shading(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{hex_color(color)}')
    pPr.append(shd)

def set_para_border_bottom(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), f'{hex_color(color)}')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_table_borders(table, color: RGBColor = BORDER_C):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), f'{hex_color(color)}')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color=TEXT, font='Calibri', underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run

def add_para(doc, text='', style=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=8):
    p = doc.add_paragraph(text, style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def add_heading(doc, text, level=1, color=WHITE, size=None, space_before=24, space_after=6):
    sizes = {1: 26, 2: 19, 3: 14, 4: 11}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if level == 2:
        set_para_border_bottom(p, BORDER_C)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(size or sizes[level])
    run.font.color.rgb = color
    return p

def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = ACCENT2
    return p

def add_body(doc, text, color=TEXT, size=10.5, space_before=0, space_after=7, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    return p

def add_code_block(doc, code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        set_para_shading(p, CODE_BG)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Cascadia Code'
        run.font.size = Pt(9)
        run.font.color.rgb = GREEN_LT
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    set_para_border_bottom(p, BORDER_C)

def add_diagram_placeholder(doc, fig_num, title, description):
    """Renders a styled diagram placeholder box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, ACCENT)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, CARD)
    cell.width = Inches(6.2)

    p_title = cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after  = Pt(4)
    r = p_title.add_run(f'[ Figure {fig_num} — {title} ]')
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT2

    p_desc = cell.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_before = Pt(0)
    p_desc.paragraph_format.space_after  = Pt(12)
    r2 = p_desc.add_run(description)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(9)
    r2.font.color.rgb = MUTED
    r2.italic = True

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(16)
    cr = cap.add_run(f'Figure {fig_num} — {title}')
    cr.italic = True
    cr.font.name = 'Calibri'
    cr.font.size = Pt(9)
    cr.font.color.rgb = MUTED

def add_callout(doc, icon, title, body, bg=SURFACE, border=ACCENT):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, border)
    # icon cell
    ic = table.rows[0].cells[0]
    ic.width = Inches(0.45)
    set_cell_bg(ic, bg)
    p = ic.add_paragraph(icon)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(0)
    p.runs[0].font.size = Pt(14)
    # body cell
    bc = table.rows[0].cells[1]
    set_cell_bg(bc, bg)
    pt = bc.add_paragraph()
    pt.paragraph_format.space_before = Pt(6)
    pt.paragraph_format.space_after  = Pt(2)
    r = pt.add_run(title)
    r.bold = True
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.color.rgb = WHITE
    pb = bc.add_paragraph(body)
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after  = Pt(6)
    pb.runs[0].font.name = 'Calibri'
    pb.runs[0].font.size = Pt(9.5)
    pb.runs[0].font.color.rgb = TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_feature_row(doc, icon, title, desc):
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table, BORDER_C)
    ic = table.rows[0].cells[0]
    ic.width = Inches(0.5)
    set_cell_bg(ic, CARD)
    p = ic.add_paragraph(icon)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    p.runs[0].font.size = Pt(16)
    bc = table.rows[0].cells[1]
    set_cell_bg(bc, CARD)
    pt = bc.add_paragraph()
    pt.paragraph_format.space_before = Pt(4)
    pt.paragraph_format.space_after  = Pt(2)
    r = pt.add_run(title)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = WHITE
    pb = bc.add_paragraph(desc)
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after  = Pt(6)
    pb.runs[0].font.name = 'Calibri'; pb.runs[0].font.size = Pt(9.5); pb.runs[0].font.color.rgb = MUTED
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_decision_card(doc, question, chosen, chosen_why, rejected_list):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(question)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = WHITE

    # chosen
    table_c = doc.add_table(rows=1, cols=1)
    set_table_borders(table_c, GREEN)
    cc = table_c.rows[0].cells[0]
    set_cell_bg(cc, RGBColor(0x14, 0x53, 0x2d))
    pc = cc.add_paragraph()
    pc.paragraph_format.space_before = Pt(4); pc.paragraph_format.space_after = Pt(4)
    rc1 = pc.add_run('CHOSEN  ')
    rc1.bold = True; rc1.font.name = 'Calibri'; rc1.font.size = Pt(9); rc1.font.color.rgb = GREEN_LT
    rc2 = pc.add_run(f'{chosen} — {chosen_why}')
    rc2.font.name = 'Calibri'; rc2.font.size = Pt(9.5); rc2.font.color.rgb = TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    for rej in rejected_list:
        table_r = doc.add_table(rows=1, cols=1)
        set_table_borders(table_r, RGBColor(0x44, 0x14, 0x1f))
        rc = table_r.rows[0].cells[0]
        set_cell_bg(rc, RGBColor(0x1f, 0x0a, 0x10))
        pr = rc.add_paragraph()
        pr.paragraph_format.space_before = Pt(4); pr.paragraph_format.space_after = Pt(4)
        rr1 = pr.add_run('REJECTED  ')
        rr1.bold = True; rr1.font.name = 'Calibri'; rr1.font.size = Pt(9); rr1.font.color.rgb = ROSE
        rr2 = pr.add_run(rej)
        rr2.font.name = 'Calibri'; rr2.font.size = Pt(9.5); rr2.font.color.rgb = MUTED
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_node_spec(doc, icon, kind_label, kind_code, subtitle, accent_color, fields):
    """fields: list of (name, type_str, description)"""
    # Header row
    ht = doc.add_table(rows=1, cols=1)
    set_table_borders(ht, accent_color)
    hc = ht.rows[0].cells[0]
    set_cell_bg(hc, CARD)
    hp = hc.add_paragraph()
    hp.paragraph_format.space_before = Pt(6); hp.paragraph_format.space_after = Pt(0)
    hp.add_run(f'{icon}  ').font.size = Pt(14)
    r1 = hp.add_run(kind_label)
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(12); r1.font.color.rgb = accent_color
    r2 = hp.add_run(f'  |  type: {kind_code}  ·  {subtitle}')
    r2.font.name = 'Cascadia Code'; r2.font.size = Pt(9); r2.font.color.rgb = MUTED

    # Fields table
    ft = doc.add_table(rows=len(fields), cols=3)
    set_table_borders(ft, BORDER_C)
    ft.columns[0].width = Inches(1.5)
    ft.columns[1].width = Inches(1.3)
    ft.columns[2].width = Inches(3.4)
    for i, (fname, ftype, fdesc) in enumerate(fields):
        row = ft.rows[i]
        set_cell_bg(row.cells[0], CARD)
        set_cell_bg(row.cells[1], CARD)
        set_cell_bg(row.cells[2], CARD)
        p0 = row.cells[0].add_paragraph(fname)
        p0.runs[0].bold = True; p0.runs[0].font.name = 'Cascadia Code'; p0.runs[0].font.size = Pt(9); p0.runs[0].font.color.rgb = WHITE
        p1 = row.cells[1].add_paragraph(ftype)
        p1.runs[0].font.name = 'Cascadia Code'; p1.runs[0].font.size = Pt(9); p1.runs[0].font.color.rgb = ACCENT2
        p2 = row.cells[2].add_paragraph(fdesc)
        p2.runs[0].font.name = 'Calibri'; p2.runs[0].font.size = Pt(9.5); p2.runs[0].font.color.rgb = MUTED
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

set_doc_bg(doc, INK)

# ── COVER ──────────────────────────────────────────────────────────────────────
p_tag = doc.add_paragraph()
p_tag.paragraph_format.space_before = Pt(40)
p_tag.paragraph_format.space_after  = Pt(20)
rt = p_tag.add_run('TREDENCE STUDIO · AI AGENTS ENGINEERING')
rt.bold = True; rt.font.name = 'Calibri'; rt.font.size = Pt(9); rt.font.color.rgb = ACCENT2

p_h1 = doc.add_paragraph()
p_h1.paragraph_format.space_before = Pt(0)
p_h1.paragraph_format.space_after  = Pt(12)
rh = p_h1.add_run('HR Workflow Designer')
rh.bold = True; rh.font.name = 'Calibri'; rh.font.size = Pt(34); rh.font.color.rgb = WHITE

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(36)
rs = p_sub.add_run(
    'Full Stack Engineering Internship — Case Study Solution Document.\n'
    'A visual drag-and-drop workflow builder for HR processes, architected with '
    'React, TypeScript, React Flow, and Zustand.'
)
rs.font.name = 'Calibri'; rs.font.size = Pt(12); rs.font.color.rgb = MUTED

# Meta table
meta = doc.add_table(rows=2, cols=3)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(meta, BORDER_C)
meta_data = [
    ('Candidate', 'Vishal'),
    ('Role', 'Full Stack Intern'),
    ('Deliverable', 'Functional Prototype'),
    ('Stack', 'React · TS · Zustand'),
    ('Date', 'April 2026'),
    ('Repo', 'github.com/root-daemon/\nhr-workflow-designer'),
]
for idx, (label, val) in enumerate(meta_data):
    r, c = divmod(idx, 3)
    cell = meta.rows[r].cells[c]
    set_cell_bg(cell, SURFACE)
    pl = cell.add_paragraph(label.upper())
    pl.paragraph_format.space_before = Pt(6); pl.paragraph_format.space_after = Pt(2)
    pl.runs[0].font.name = 'Calibri'; pl.runs[0].font.size = Pt(8); pl.runs[0].font.color.rgb = MUTED; pl.runs[0].bold = True
    pv = cell.add_paragraph(val)
    pv.paragraph_format.space_before = Pt(0); pv.paragraph_format.space_after = Pt(6)
    pv.runs[0].font.name = 'Calibri'; pv.runs[0].font.size = Pt(10); pv.runs[0].font.color.rgb = WHITE; pv.runs[0].bold = True

doc.add_page_break()

# ── TOC ────────────────────────────────────────────────────────────────────────
add_heading(doc, 'Table of Contents', level=2, space_before=8)
toc_items = [
    ('01', 'Executive Summary'),
    ('02', 'System Architecture'),
    ('03', 'Component Architecture'),
    ('04', 'State Management'),
    ('05', 'Data Flow & User Interactions'),
    ('06', 'Node Type Specifications'),
    ('07', 'Mock API Layer'),
    ('08', 'Simulation Engine'),
    ('09', 'Key Design Decisions'),
    ('10', 'Scalability & Extension Points'),
    ('11', 'Feature Delivery Summary'),
    ('12', 'What I\'d Build Next'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    set_para_border_bottom(p, BORDER_C)
    rn = p.add_run(f'{num}  ')
    rn.bold = True; rn.font.name = 'Cascadia Code'; rn.font.size = Pt(9); rn.font.color.rgb = ACCENT2
    rt2 = p.add_run(title)
    rt2.font.name = 'Calibri'; rt2.font.size = Pt(11); rt2.font.color.rgb = TEXT

doc.add_page_break()

# ══════════════════════════════════════════════════════ §1 EXECUTIVE SUMMARY ══
add_label(doc, 'Section 01')
add_heading(doc, 'Executive Summary', level=1)
add_body(doc,
    'A working prototype of an HR Workflow Designer, delivered within the 4–6 hour '
    'time-box. The application satisfies every mandatory requirement and implements '
    'four of the six bonus features.', space_after=16)

features = [
    ('🎨', 'Visual Canvas',
     'Drag-and-drop React Flow canvas with 5 typed node varieties, live edge connections, and a minimap.'),
    ('⚙️', 'Node Config Forms',
     'Per-node dynamic forms with controlled components, validation, and type-safe state propagation via Zustand.'),
    ('🔌', 'Mock API Layer',
     'Pure async functions simulating GET /automations and POST /simulate with realistic latency.'),
    ('⚗️', 'Simulation Sandbox',
     'Step-by-step animated execution log, graph validation, topological sort, and raw JSON inspector.'),
    ('↩️', 'Undo / Redo',
     '50-step snapshot history baked into the Zustand store. Every destructive action is reversible.'),
    ('📦', 'Import / Export',
     'Full workflow graph serialisation to JSON — downloadable or clipboard — and importable back in.'),
]
for icon, title, desc in features:
    add_feature_row(doc, icon, title, desc)

add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §2 SYSTEM ARCHITECTURE ═
add_label(doc, 'Section 02')
add_heading(doc, 'System Architecture', level=1)
add_body(doc,
    'The application is a pure client-side SPA. There is no backend — the API layer '
    'is simulated with local async functions. This satisfies the "no auth or backend '
    'persistence required" constraint while keeping the mock-to-real migration path trivial.',
    space_after=16)

add_diagram_placeholder(doc, 1, 'High-Level System Architecture',
    'Presentation Layer → State Layer (Zustand) → Logic Layer → Mock API Layer\n'
    'All layers run in the browser; no network calls are made.')

add_heading(doc, 'Layer Responsibilities', level=3)

layer_table = doc.add_table(rows=6, cols=3)
layer_table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(layer_table, BORDER_C)
headers = ['Layer', 'Module(s)', 'Responsibility']
h_colors = [ACCENT2, ACCENT2, ACCENT2]
for ci, (h, hc) in enumerate(zip(headers, h_colors)):
    cell = layer_table.rows[0].cells[ci]
    set_cell_bg(cell, CARD)
    p = cell.add_paragraph(h.upper())
    p.runs[0].bold = True; p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = MUTED
rows_data = [
    ('Presentation', 'components/**', 'Render UI, forward events to the store'),
    ('State', 'store/workflowStore.ts', 'Single source of truth for all mutable application state'),
    ('Logic', 'api/mockApi.ts (internal)', 'Graph validation, topological sort, step generation'),
    ('API', 'api/mockApi.ts (exports)', 'Async contract — swap for real fetch without changing callers'),
    ('Types', 'types/workflow.ts', 'Shared TypeScript interfaces referenced by all layers'),
]
for ri, (l, m, resp) in enumerate(rows_data):
    row = layer_table.rows[ri + 1]
    for ci2, (val, col) in enumerate([(l, ACCENT2), (m, ACCENT2), (resp, TEXT)]):
        set_cell_bg(row.cells[ci2], SURFACE if ri % 2 == 0 else INK)
        p = row.cells[ci2].add_paragraph(val)
        p.runs[0].font.name = 'Calibri' if ci2 != 1 else 'Cascadia Code'
        p.runs[0].font.size = Pt(9.5 if ci2 != 1 else 9)
        p.runs[0].font.color.rgb = col if ci2 != 2 else TEXT
doc.add_paragraph().paragraph_format.space_after = Pt(8)

add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §3 COMPONENT ARCH ══════
add_label(doc, 'Section 03')
add_heading(doc, 'Component Architecture', level=1)
add_body(doc,
    'Components are organised by domain, not by type. Each directory owns a single '
    'concern: canvas wiring, node rendering, form editing, or sandbox simulation.',
    space_after=16)

add_diagram_placeholder(doc, 2, 'React Component Tree',
    'App.tsx → Toolbar / NodePalette / WorkflowCanvas / NodeFormPanel / SandboxPanel\n'
    'WorkflowCanvas → StartNode / TaskNode / ApprovalNode / AutomatedNode / EndNode → BaseNode\n'
    'NodeFormPanel → StartNodeForm / TaskNodeForm / ApprovalNodeForm / AutomatedNodeForm / EndNodeForm → KeyValueEditor')

add_heading(doc, 'Folder Structure', level=3)
add_code_block(doc, """src/
├── api/
│   └── mockApi.ts            ← async API contract + mock data
├── components/
│   ├── canvas/
│   │   ├── NodePalette.tsx   ← draggable node type list (left sidebar)
│   │   └── WorkflowCanvas.tsx← ReactFlow + drop handler
│   ├── forms/
│   │   ├── NodeFormPanel.tsx ← kind-aware form dispatcher
│   │   ├── StartNodeForm.tsx
│   │   ├── TaskNodeForm.tsx
│   │   ├── ApprovalNodeForm.tsx
│   │   ├── AutomatedNodeForm.tsx
│   │   ├── EndNodeForm.tsx
│   │   └── KeyValueEditor.tsx← shared kv-pair component
│   ├── nodes/
│   │   ├── BaseNode.tsx      ← shared node shell
│   │   ├── {Start,Task,Approval,Automated,End}Node.tsx
│   │   └── nodeStyles.ts     ← colour / icon registry
│   ├── sandbox/
│   │   └── SandboxPanel.tsx  ← simulate + timeline + JSON viewer
│   └── ui/
│       └── Toolbar.tsx       ← top bar actions
├── store/
│   └── workflowStore.ts      ← Zustand store (all state)
└── types/
    └── workflow.ts           ← shared TS interfaces""")

add_callout(doc, '💡', 'Design principle',
    'Every directory is independently deletable. Removing sandbox/ does not break canvas or forms. '
    'Adding a new node type only touches nodes/, forms/, types/workflow.ts, and one switch in NodeFormPanel.')

add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §4 STATE MANAGEMENT ════
add_label(doc, 'Section 04')
add_heading(doc, 'State Management', level=1)
add_body(doc,
    'All mutable application state lives in a single Zustand store. Components subscribe '
    'with fine-grained selectors, so a form edit never re-renders the canvas and vice-versa.',
    space_after=16)

add_diagram_placeholder(doc, 3, 'Zustand Store State Machine',
    'Idle → NodeAdded / NodeDeleted / EdgeConnected / ConfigEdited / NodeSelected / Undone / Redone\n'
    'Every mutating action optionally snapshots history before applying.')

add_heading(doc, 'Store Shape', level=3)
add_code_block(doc, """interface WorkflowState {
  // Core data
  nodes:          WorkflowNode[]
  edges:          WorkflowEdge[]
  selectedNodeId: string | null

  // History (undo/redo)
  history:      { nodes: WorkflowNode[]; edges: WorkflowEdge[] }[]
  historyIndex: number           // pointer into history ring

  // React Flow handlers
  onNodesChange:  (changes: NodeChange[]) => void
  onEdgesChange:  (changes: EdgeChange[]) => void
  onConnect:      (connection: Connection) => void

  // Domain actions
  addNode:          (kind: NodeKind, position: XYPosition) => void
  deleteNode:       (id: string) => void
  deleteEdge:       (id: string) => void
  updateNodeConfig: (id: string, patch: Partial<NodeConfig>) => void
  selectNode:       (id: string | null) => void
  pushHistory:      () => void
  undo:             () => void
  redo:             () => void

  // Serialisation
  exportWorkflow: () => string
  importWorkflow: (json: string) => void
}""")

add_callout(doc, '✓', 'Undo / Redo implementation',
    'pushHistory() deep-clones the current { nodes, edges } before every mutating action. '
    'The ring keeps the last 50 snapshots. undo() decrements the pointer and restores; '
    'redo() increments it. No external library needed.',
    bg=RGBColor(0x05, 0x2e, 0x16), border=GREEN)

add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §5 DATA FLOW ═══════════
add_label(doc, 'Section 05')
add_heading(doc, 'Data Flow & User Interactions', level=1)
add_body(doc,
    'Two primary interaction flows cover the full lifecycle of a workflow: '
    'authoring (drag → configure → connect) and testing (simulate → inspect).',
    space_after=16)

add_heading(doc, 'Authoring Flow', level=3)
add_diagram_placeholder(doc, 4, 'Authoring Sequence Diagram',
    'HR Admin → NodePalette (dragstart) → WorkflowCanvas (drop) → Zustand addNode\n'
    '→ click node → selectNode → NodeFormPanel opens → edit fields → updateNodeConfig\n'
    '→ drag handle → onConnect → addEdge (animated, styled indigo)')

add_heading(doc, 'Simulation Flow', level=3)
add_diagram_placeholder(doc, 5, 'Simulation Sequence Diagram',
    'HR Admin → SandboxPanel "Run" → simulateWorkflow({ nodes, edges })\n'
    '→ Graph Validator → [errors → show red callout] or [pass → topologicalSort]\n'
    '→ generate SimulateStep per node → animate timeline → show detail panel')

add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §6 NODE SPECS ══════════
add_label(doc, 'Section 06')
add_heading(doc, 'Node Type Specifications', level=1)
add_body(doc,
    'Five node types cover the complete HR workflow lifecycle. Each has a unique visual '
    'identity, config shape, and set of constraints.',
    space_after=16)

add_diagram_placeholder(doc, 6, 'Node Connection Topology',
    'START → TASK / APPROVAL / AUTOMATED\n'
    'TASK → APPROVAL / AUTOMATED / END\n'
    'APPROVAL → TASK / AUTOMATED / END\n'
    'AUTOMATED → TASK / END\n'
    'Start has no target handle · End has no source handle')

add_node_spec(doc, '▶', 'Start Node', 'start', 'No target handle · One per workflow',
    GREEN, [
        ('title', 'string', 'Human-readable workflow name'),
        ('metadata', 'KeyValue[]', 'Arbitrary key-value pairs for workflow context (e.g. department: HR)'),
    ])

add_node_spec(doc, '✓', 'Task Node', 'task', 'Represents a human action to be completed',
    BLUE, [
        ('title', 'string *', 'Required. Task name shown on node and in simulation log'),
        ('description', 'string', 'Detailed instructions for the assignee'),
        ('assignee', 'string', 'Name or email of responsible person'),
        ('dueDate', 'string', 'ISO date string shown on the node card'),
        ('customFields', 'KeyValue[]', 'Extensible fields (e.g. priority: High)'),
    ])

add_node_spec(doc, '◆', 'Approval Node', 'approval', 'Blocks until a role-holder approves',
    AMBER, [
        ('title', 'string', 'Approval step label'),
        ('approverRole', 'enum', 'Manager · HRBP · Director · VP · C-Suite · Legal'),
        ('autoApproveThreshold', 'number', 'Days before auto-approval triggers. 0 = disabled.'),
    ])

add_node_spec(doc, '⚡', 'Automated Step Node', 'automated', 'System-triggered action fetched from API',
    PURPLE, [
        ('title', 'string', 'Step label shown on node'),
        ('actionId', 'string', 'ID from GET /automations (e.g. send_email)'),
        ('actionParams', 'Record<string,string>', 'Dynamic fields keyed by action.params[] — rendered at runtime from API response'),
    ])

add_node_spec(doc, '■', 'End Node', 'end', 'No source handle · Terminal state',
    ROSE, [
        ('endMessage', 'string', 'Completion message shown in summary'),
        ('summaryFlag', 'boolean', 'Toggle — whether to compile a workflow execution summary report'),
    ])

add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §7 MOCK API ═════════════
add_label(doc, 'Section 07')
add_heading(doc, 'Mock API Layer', level=1)
add_body(doc,
    'The API contract is expressed as typed async functions in src/api/mockApi.ts. '
    'They introduce realistic latency via setTimeout and return data that matches the '
    'shapes the case study specifies. Swapping them for real fetch calls requires '
    'zero changes to call sites.',
    space_after=16)

add_diagram_placeholder(doc, 7, 'API Layer Call Graph',
    'AutomatedNodeForm → getAutomations() → MOCK_AUTOMATIONS → AutomationAction[]\n'
    'SandboxPanel → simulateWorkflow(req) → validateGraph() + topologicalSort() → SimulateResult')

add_heading(doc, 'Contract — GET /automations', level=3)
add_code_block(doc, """interface AutomationAction {
  id:     string    // "send_email"
  label:  string    // "Send Email"
  params: string[]  // ["to", "subject", "body"]
}

// 6 mock actions provided
[
  { id: "send_email",       label: "Send Email",            params: ["to","subject","body"] },
  { id: "generate_doc",    label: "Generate Document",     params: ["template","recipient"] },
  { id: "create_ticket",   label: "Create JIRA Ticket",    params: ["project","summary","priority"] },
  { id: "slack_notify",    label: "Send Slack Message",    params: ["channel","message"] },
  { id: "create_account",  label: "Create System Account", params: ["username","role","department"] },
  { id: "schedule_meeting",label: "Schedule Meeting",      params: ["attendees","title","duration"] },
]""")

add_heading(doc, 'Contract — POST /simulate', level=3)
add_code_block(doc, """interface SimulateRequest  { nodes: WorkflowNode[]; edges: WorkflowEdge[] }

interface SimulateResult {
  success:         boolean
  steps:           SimulateStep[]
  errors:          string[]
  totalDurationMs: number
}

interface SimulateStep {
  nodeId:     string
  label:      string
  kind:       NodeKind
  status:     "success" | "error" | "skipped"
  message:    string
  durationMs: number
}""")

add_callout(doc, '⚠️', 'Production migration path',
    'Replace the two exported functions in mockApi.ts with '
    'fetch("/api/automations") and fetch("/api/simulate", { method: "POST", body: JSON.stringify(req) }). '
    'The rest of the codebase is untouched.',
    bg=RGBColor(0x27, 0x1a, 0x02), border=AMBER)

add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §8 SIMULATION ENGINE ════
add_label(doc, 'Section 08')
add_heading(doc, 'Simulation Engine', level=1)
add_body(doc,
    'The simulation runs in three phases: validate the graph structure, sort nodes into '
    'execution order via a topological BFS (Kahn\'s algorithm), then generate a '
    'SimulateStep for each node.',
    space_after=16)

add_diagram_placeholder(doc, 8, 'Simulation Engine Flowchart',
    'START → Check ≥1 Start node → Check ≤1 Start node → Check ≥1 End node\n'
    '→ Check no disconnected nodes → [errors? → return failure]\n'
    '→ Build adjacency + in-degree map → BFS queue from zero-in-degree nodes\n'
    '→ Dequeue → append → decrement neighbors → [cycle? → fallback to original order]\n'
    '→ Generate SimulateStep per node → sum durations → return SimulateResult')

add_heading(doc, "Kahn's Algorithm — Implementation", level=3)
add_code_block(doc, """function topologicalSort(nodes, edges) {
  const adj      = {}   // nodeId → neighbor ids[]
  const inDegree = {}   // nodeId → incoming edge count

  nodes.forEach(n => { adj[n.id] = []; inDegree[n.id] = 0 })
  edges.forEach(e => { adj[e.source].push(e.target); inDegree[e.target]++ })

  // Start BFS from nodes with no incoming edges
  const queue  = nodes.filter(n => inDegree[n.id] === 0)
  const sorted = []

  while (queue.length > 0) {
    const node = queue.shift()
    sorted.push(node)
    for (const neighborId of adj[node.id]) {
      inDegree[neighborId]--
      if (inDegree[neighborId] === 0)
        queue.push(nodes.find(n => n.id === neighborId))
    }
  }

  // Cycle guard: if sorted < nodes a cycle exists → original order
  return sorted.length === nodes.length ? sorted : nodes
}""")

add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §9 DESIGN DECISIONS ════
add_label(doc, 'Section 09')
add_heading(doc, 'Key Design Decisions', level=1)
add_body(doc,
    'Each major architectural choice had at least one credible alternative. '
    'Here\'s what was chosen, what was rejected, and why.',
    space_after=16)

decisions = [
    ('State management library',
     'Zustand',
     'single flat store, actions co-located, fine-grained selectors, zero boilerplate, built-in snapshot for undo/redo',
     [
         'Redux Toolkit — powerful but excessive ceremony for a single-page workflow tool; adds 3× the boilerplate without meaningful benefit at this scale',
         'React Context + useReducer — viable, but causes entire subtrees to re-render on any state change unless carefully memoized; Zustand handles this transparently',
     ]),
    ('API mocking strategy',
     'Pure async functions in mockApi.ts',
     'zero dependencies, trivially replaceable with real fetch, deterministic, no service worker setup friction',
     [
         'MSW (Mock Service Worker) — excellent for integration testing but adds service worker complexity not warranted here',
         'json-server — requires a separate process, breaking the single "npm run dev" developer experience',
     ]),
    ('Node form architecture',
     'Separate form component per node type (open/closed principle)',
     'NodeFormPanel dispatches on kind; adding a new type adds one file and changes one switch case',
     [
         'Single generic form driven by a schema object — fewer files but harder to maintain as field types diverge; toggle vs select vs kv-list don\'t share a natural base component',
     ]),
    ('Node style system',
     'Central nodeStyles.ts registry',
     'mapping each NodeKind to gradient/border/icon/badge classes — palette, nodes, and forms all import from one place',
     [
         'Inline styles per component — duplicates colour decisions, breaks consistency when a colour changes',
     ]),
    ('Undo / Redo mechanism',
     'Deep-clone snapshot ring inside Zustand (50 entries)',
     'straightforward, testable, no dependency, works for any state shape',
     [
         'use-undoable / immer-with-patches — more memory-efficient (patches vs full snapshots) but adds complexity; 50 snapshots of a typical graph is ≈200 kB max',
     ]),
    ('Styling system',
     'Tailwind CSS v3 with custom colour palette in tailwind.config.js',
     'semantic tokens (bg-canvas, text-muted) keep components readable without context-switching to CSS files',
     [
         'CSS Modules — valid but requires context-switching between JS and CSS files; Tailwind keeps styling co-located',
     ]),
]
for q, chosen, why, rejected in decisions:
    add_decision_card(doc, q, chosen, why, rejected)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §10 SCALABILITY ═════════
add_label(doc, 'Section 10')
add_heading(doc, 'Scalability & Extension Points', level=1)
add_body(doc,
    'The architecture is designed to grow. Below are the primary extension points '
    'and how they\'d be exercised in a production context.',
    space_after=16)

add_diagram_placeholder(doc, 9, 'Prototype → Production Migration Path',
    'mockApi.ts async functions  →  FastAPI backend\n'
    'Zustand in-memory           →  PostgreSQL (workflow table) + Firestore (real-time sync)\n'
    'Vite SPA                    →  Kubernetes deployment + Azure OIDC auth')

add_heading(doc, 'Adding a New Node Type (3 Steps)', level=3)
add_code_block(doc, """// 1. Add to NodeKind union in types/workflow.ts
type NodeKind = 'start' | 'task' | 'approval' | 'automated' | 'end' | 'condition'

// 2. Register visual identity in nodeStyles.ts
NODE_STYLES['condition'] = {
  gradient: 'from-cyan-500/20 to-cyan-600/10',
  border:   'border-cyan-500/50',
  icon:     '?',
  badge:    'bg-cyan-500/20 text-cyan-400',
  badgeText: 'CONDITION',
}

// 3. Create ConditionNode.tsx, ConditionNodeForm.tsx,
//    add defaultConfig() case, add <ConditionNodeForm /> in NodeFormPanel""")

add_heading(doc, 'Performance Considerations', level=4)
add_body(doc,
    'All node components are wrapped in React.memo. React Flow only re-renders nodes '
    'whose position or data reference changes. The Zustand selector pattern '
    '(s => s.selectedNodeId) means the form panel re-renders in isolation from the canvas '
    'on selection changes.')

add_heading(doc, 'Type Safety', level=4)
add_body(doc,
    'Every node\'s config is typed as a discriminated union via NodeKind. '
    'defaultConfig(kind) uses a satisfies assertion so the TS compiler verifies each '
    'branch returns the correct shape. Adding a field to TaskConfig immediately surfaces '
    'as a type error in the form.')

add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §11 DELIVERY SUMMARY ════
add_label(doc, 'Section 11')
add_heading(doc, 'Feature Delivery Summary', level=1)
add_body(doc, 'Assessment criteria mapped against delivered implementation.', space_after=16)

delivery = [
    ('Drag-and-drop canvas',           '✓ Done',     'HTML5 drag from palette → screenToFlowPosition → addNode()'),
    ('5 custom node types',            '✓ Done',     'Start · Task · Approval · Automated · End — each React.memo component'),
    ('Connect nodes with edges',       '✓ Done',     'Animated edges via onConnect → Zustand addEdge'),
    ('Select node to edit',            '✓ Done',     'Click → selectNode(id) → right panel slides in'),
    ('Delete nodes / edges',           '✓ Done',     'Node: hover delete btn or form panel button · Edge: click → confirm dialog'),
    ('Node configuration forms',       '✓ Done',     'Controlled components, real-time preview on node card, full type safety'),
    ('Dynamic AutomatedNode params',   '✓ Done',     'Fetches action list on mount, renders param inputs from action.params[]'),
    ('Key-value custom fields',        '✓ Done',     'KeyValueEditor shared between Start and Task forms'),
    ('GET /automations',               '✓ Done',     '6 mock actions, 300 ms artificial latency, loading state shown'),
    ('POST /simulate',                 '✓ Done',     'Validates → topological sort → generates steps with realistic timing'),
    ('Workflow test / sandbox panel',  '✓ Done',     'Slide-up overlay, animated step timeline, detail sidebar, JSON tab'),
    ('Graph validation',               '✓ Done',     'Missing Start/End, multiple Start nodes, disconnected nodes — all caught'),
    ('Bonus: Undo / Redo',             '✓ Done',     '50-step snapshot ring in Zustand store'),
    ('Bonus: Export / Import JSON',    '✓ Done',     'Download file + clipboard copy + import from file picker'),
    ('Bonus: Minimap + Zoom controls', '✓ Done',     'React Flow built-ins with dark-theme overrides'),
    ('Bonus: Auto-validate constraints','✓ Done',    'Validated at simulation time with human-readable error messages'),
    ('Bonus: Node templates',          'Not done',   'Design is straightforward — predefined JSON arrays via existing importWorkflow()'),
    ('Bonus: Auto-layout',             'Not done',   'Would use Dagre or ELK; not critical within time-box'),
    ('Bonus: Node version history',    'Not done',   'Would store per-node config changelog in the store; post time-box feature'),
]

dt = doc.add_table(rows=len(delivery)+1, cols=3)
dt.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(dt, BORDER_C)
for ci2, h in enumerate(['Requirement', 'Status', 'Implementation Detail']):
    cell = dt.rows[0].cells[ci2]
    set_cell_bg(cell, CARD)
    p = cell.add_paragraph(h.upper())
    p.runs[0].bold = True; p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = MUTED

for ri2, (req, status, detail) in enumerate(delivery):
    row = dt.rows[ri2 + 1]
    bg = SURFACE if ri2 % 2 == 0 else INK
    for ci3 in range(3): set_cell_bg(row.cells[ci3], bg)
    p0 = row.cells[0].add_paragraph(req)
    p0.runs[0].font.name = 'Calibri'; p0.runs[0].font.size = Pt(9.5); p0.runs[0].font.color.rgb = TEXT
    status_color = GREEN if status == '✓ Done' else AMBER
    p1 = row.cells[1].add_paragraph(status)
    p1.runs[0].font.name = 'Calibri'; p1.runs[0].font.size = Pt(9.5)
    p1.runs[0].font.color.rgb = status_color; p1.runs[0].bold = True
    p2 = row.cells[2].add_paragraph(detail)
    p2.runs[0].font.name = 'Calibri'; p2.runs[0].font.size = Pt(9); p2.runs[0].font.color.rgb = MUTED

doc.add_paragraph().paragraph_format.space_after = Pt(8)
add_divider(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════ §12 WHAT NEXT ══════════
add_label(doc, 'Section 12')
add_heading(doc, "What I'd Build Next", level=1)
add_body(doc,
    'Prioritised by impact and approximate complexity.',
    space_after=16)

next_items = [
    ('🔴', 'HIGH', 'Visual validation errors on nodes',
     'Red outline + tooltip on invalid nodes (e.g. Automated node with no action selected, Task with '
     'no title). Surface errors before simulation, not during.'),
    ('🔴', 'HIGH', 'Keyboard shortcuts',
     'Ctrl+Z / Ctrl+Y for undo/redo (global keydown listener → store actions), Ctrl+D to duplicate a node, '
     'Del to delete selected node without hover.'),
    ('🟠', 'MED', 'localStorage autosave',
     'Persist workflow state on every store mutation with a Zustand middleware. Restore on page load with '
     'a "continue editing" banner. Zero backend required.'),
    ('🟠', 'MED', 'Workflow templates',
     'Pre-built JSON payloads for Onboarding, Leave Approval, and Document Verification — loaded via the '
     'existing importWorkflow() path.'),
    ('🟡', 'MED', 'Dagre auto-layout',
     'Single button to re-layout the graph using dagre. Critical for imported workflows with random positions. '
     'React Flow\'s setNodes accepts new positions directly.'),
    ('🟡', 'MED', 'Playwright E2E test suite',
     'Cover: drag-drop a node, configure it, connect two nodes, run simulation, check step count, '
     'export JSON, import it back. Verifies the golden path stays green.'),
    ('🔵', 'PROD', 'FastAPI persistence',
     'Swap mockApi.ts for real endpoints. Store workflows in PostgreSQL (workflow table with JSONB graph column). '
     'Add versioning via workflow_versions.'),
    ('🔵', 'PROD', 'Azure OIDC authentication',
     'Add msal-react for Azure AD login. Pass JWT in Authorization header. Scope workflows to authenticated tenants.'),
    ('🔵', 'PROD', 'Real-time collaborative editing',
     'WebSocket or Firestore-backed live sync — multiple HR admins see each other\'s edits. '
     'Conflict resolution via Y.js CRDT.'),
]

for icon, priority, title, desc in next_items:
    t = doc.add_table(rows=1, cols=2)
    set_table_borders(t, BORDER_C)
    ic = t.rows[0].cells[0]
    ic.width = Inches(0.55)
    bc2 = t.rows[0].cells[1]
    pri_colors = {'HIGH': ROSE, 'MED': AMBER, 'PROD': ACCENT2}
    pri_bgs = {'HIGH': RGBColor(0x1a, 0x05, 0x08), 'MED': RGBColor(0x1a, 0x11, 0x03), 'PROD': RGBColor(0x0e, 0x0f, 0x26)}
    set_cell_bg(ic, pri_bgs.get(priority, SURFACE))
    set_cell_bg(bc2, SURFACE)
    pi = ic.add_paragraph(icon)
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.paragraph_format.space_before = Pt(6); pi.paragraph_format.space_after = Pt(2)
    pi.runs[0].font.size = Pt(16)
    pp = ic.add_paragraph(priority)
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(0); pp.paragraph_format.space_after = Pt(6)
    pp.runs[0].font.name = 'Calibri'; pp.runs[0].font.size = Pt(8)
    pp.runs[0].font.color.rgb = pri_colors.get(priority, MUTED); pp.runs[0].bold = True
    pt2 = bc2.add_paragraph()
    pt2.paragraph_format.space_before = Pt(6); pt2.paragraph_format.space_after = Pt(3)
    rt3 = pt2.add_run(title)
    rt3.bold = True; rt3.font.name = 'Calibri'; rt3.font.size = Pt(11); rt3.font.color.rgb = WHITE
    pd = bc2.add_paragraph(desc)
    pd.paragraph_format.space_before = Pt(0); pd.paragraph_format.space_after = Pt(6)
    pd.runs[0].font.name = 'Calibri'; pd.runs[0].font.size = Pt(9.5); pd.runs[0].font.color.rgb = MUTED
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── FOOTER ─────────────────────────────────────────────────────────────────────
add_divider(doc)
p_foot = doc.add_paragraph()
p_foot.paragraph_format.space_before = Pt(8)
p_foot.paragraph_format.space_after  = Pt(4)
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf1 = p_foot.add_run('HR Workflow Designer')
rf1.bold = True; rf1.font.name = 'Calibri'; rf1.font.size = Pt(10); rf1.font.color.rgb = WHITE
rf2 = p_foot.add_run(' — Tredence Studio Full Stack Internship Case Study')
rf2.font.name = 'Calibri'; rf2.font.size = Pt(10); rf2.font.color.rgb = MUTED

p_foot2 = doc.add_paragraph()
p_foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf3 = p_foot2.add_run('React · TypeScript · React Flow · Zustand · Tailwind CSS  ·  ')
rf3.font.name = 'Calibri'; rf3.font.size = Pt(9); rf3.font.color.rgb = MUTED
rf4 = p_foot2.add_run('github.com/root-daemon/hr-workflow-designer')
rf4.font.name = 'Calibri'; rf4.font.size = Pt(9); rf4.font.color.rgb = ACCENT2

# ── Save ────────────────────────────────────────────────────────────────────────
out = '/Users/root-daemon/Programming/tredance/hr-workflow-designer/CASE_STUDY.docx'
doc.save(out)
print(f'Saved → {out}')
